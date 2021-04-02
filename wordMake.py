# import mmap
# import contextlib
# import time
import sys
import os
import datetime
import json
import docx as dx
from docx import Document
from docx.shared import Inches
from docx.shared import Cm


def testWith():

    # with open("test.dat", "w") as f:
    #     f.write('\x00' * 1024)

    # with open('test.dat', 'r+') as f:
    #     with contextlib.closing(mmap.mmap(f.fileno(), 1024, access=mmap.ACCESS_WRITE)) as m:
    #         for i in range(1, 10001):
    #             m.seek(0)
    #             s = "msg " + str(i)
    #             s.rjust(1024, '\x00')
    #             m.write(s)
    #             m.flush()
    #             time.sleep(1)
    count = 0
    while (True):
        print ('The count is:', count)
        count += 1
        time.sleep(1)

def jsonParsing(jsonStr):
    with open("test.txt","w") as f:
        f.write(jsonStr)

def readFromText():
    jsonStr = ""
    with open("testJson.txt","r") as f:
        jsonStr = f.read()
    return jsonStr

class WordMakeObject:
    def __init__(self):
        timeStr = datetime.datetime.now().strftime('%Y%m%d%H%M%S')
        timeStr = timeStr+".docx"
        self.fileName = timeStr

    def createDocument(self, dicItem):
        add = dicItem.get("Add")
        fileName = dicItem.get("FileName")
        if fileName != None:
            self.fileName = fileName
        if add == None or add == False or os.path.isfile(self.fileName) == False:
            self.document = Document()
        elif add == True:
            self.document = Document(self.fileName)

    def setTitle(self, dicItem):
        titleName = dicItem.get("Title")
        typeName =  dicItem.get("Type")
        if titleName != None and typeName != None:
            self.document.add_heading(titleName, typeName)
    
    def setLineText(self, dicItem):
        lineText = dicItem.get("Paragraph")
        lineTextBold =  dicItem.get("Bold")
        lineTextItalic =  dicItem.get("Italic")
        print(lineText, lineTextBold, lineTextItalic)
        if lineText != None and lineTextBold != None and lineTextItalic != None:
            p = self.document.add_paragraph(lineText)
            p.bold = lineTextBold
            p.italic = lineTextItalic
    
    def setTable(self, dicItem):
        tableList = dicItem.get("Table")
        if tableList == None:
            return
        numList = len(tableList)
        col = 0
        if numList > 0:
            col = len(tableList[0])
        table = self.document.add_table(rows=1, cols=col)
        rowCells = table.rows[0].cells
        for index in range(len(tableList[0])):
            rowCells[index].text = tableList[0][index]

        tableList.pop(0)

        for cellRow in tableList:
            rowCells = table.add_row().cells
            for index in range(len(cellRow)):
                val = cellRow[index]
                if type(val) != str:
                    val = str(val)
                rowCells[index].text = val
    
    def setPixMap(self, dicItem):
        pixMapPath = dicItem.get("PixMapName")
        pixInches = dicItem.get("Inches")
        if pixMapPath == None or pixInches == None:
            return
        self.document.add_picture(pixMapPath, width=Inches(pixInches))
        self.alignCenter()

    def alignCenter(self):
        lastParagraph = self.document.paragraphs[-1]
        lastParagraph.alignment = dx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    
    def setAddPageBreak(self, dicItem):
        pageBreak = dicItem.get("PageBreak")
        if pageBreak == None:
            return
        for i in range(pageBreak):
            self.document.add_page_break()

    def makeWord(self, dicInfo):
        numDic = len(dicInfo)
        if numDic > 0:
            self.createDocument(dicInfo[0])
        for dicItem in dicInfo:
            self.setTitle(dicItem)
            self.setLineText(dicItem)
            self.setTable(dicItem)
            self.setPixMap(dicItem)
            self.setAddPageBreak(dicItem)
        self.document.save(self.fileName)

if __name__ == '__main__':
    jsonStr = sys.argv[1]
    # jsonStr = readFromText()
    dicInfo = json.loads(jsonStr)
    makeWordObject = WordMakeObject()
    makeWordObject.makeWord(dicInfo)
    # jsonParsing(jsonStr)